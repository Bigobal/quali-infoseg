import tkinter as tk
from tkinter import filedialog, messagebox
import docx
from docx.shared import Inches
import base64
from PIL import Image
import io
import os

def ler_dados_txt(caminho_txt):
    dados = {}
    with open(caminho_txt, 'r', encoding='utf-8') as file:
        for linha in file:
            if ':' in linha:
                chave, valor = linha.strip().split(':', 1)
                dados[chave.strip()] = valor.strip()
    return dados

def substituir_texto(doc, dados):
    for paragrafo in doc.paragraphs:
        for chave, valor in dados.items():
            if chave != 'Imagem': 
                paragrafo.text = paragrafo.text.replace(chave, valor)
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for chave, valor in dados.items():
                        if chave != 'Imagem':
                            paragrafo.text = paragrafo.text.replace(chave, valor)

def substituir_imagem(doc, imagem_base64):
    try:
        imagem_decodificada = base64.b64decode(imagem_base64)
        imagem = Image.open(io.BytesIO(imagem_decodificada))  

        caminho_imagem = 'temp_imagem.png'
        imagem.save(caminho_imagem)

        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    if '{Imagem}' in celula.text:
                        celula.text = ''
                        celula.paragraphs[0].add_run().add_picture(caminho_imagem, width=Inches(1.5))

    except base64.binascii.Error:
        messagebox.showerror("Erro", "A imagem fornecida não é um Base64 válido.")
    except IOError:
        messagebox.showerror("Erro", "A imagem decodificada não é válida.")
    except Exception as e:
        messagebox.showerror("Erro", f"Ocorreu um erro ao processar a imagem: {str(e)}")

def modificar_documento(caminho_txt):
    caminho_base = os.path.dirname(os.path.abspath(__file__))
    caminho_docx = os.path.join(caminho_base, "qualificação geral.docx")

    dados = ler_dados_txt(caminho_txt)

    doc = docx.Document(caminho_docx)

    substituir_texto(doc, dados)

    if 'Imagem' in dados:
        substituir_imagem(doc, dados['Imagem'])

    nome_arquivo = dados.get('NNNN', 'documento_modificado').strip()
    nome_arquivo = nome_arquivo.replace('/', '-')
    caminho_docx_modificado = f"{nome_arquivo}.docx"

    caminho_saida = os.path.join(caminho_base, f"{nome_arquivo}.docx")
    doc.save(caminho_saida)


    messagebox.showinfo("Sucesso", f"Documento gerado salvo em: {caminho_docx_modificado}")


def selecionar_arquivo_dados():
    caminho_txt = filedialog.askopenfilename(title="Selecione o arquivo de dados", filetypes=[("Text files", "*.txt")])
    entry_dados.insert(0, caminho_txt)

def executar_modificacao():
    caminho_txt = entry_dados.get()
    if caminho_txt:
        modificar_documento(caminho_txt)
    else:
        messagebox.showwarning("Aviso", "Por favor, selecione o arquivo de dados.")

root = tk.Tk()
root.title("Qualificador Infoseg")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

label_dados = tk.Label(frame, text="Arquivo de dados:")
label_dados.grid(row=0, column=0, padx=5, pady=5)
entry_dados = tk.Entry(frame, width=50)
entry_dados.grid(row=0, column=1, padx=5, pady=5)
button_dados = tk.Button(frame, text="Selecionar", command=selecionar_arquivo_dados)
button_dados.grid(row=0, column=2, padx=5, pady=5)

button_executar = tk.Button(root, text="Gerar Documento", command=executar_modificacao)
button_executar.pack(pady=20)

root.mainloop()
